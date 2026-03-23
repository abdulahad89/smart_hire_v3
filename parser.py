import fitz  # PyMuPDF
import docx
import re
from typing import Dict, List, Any
from pathlib import Path
import streamlit as st

class EnhancedResumeParser:
    """Enhanced resume parser with skill extraction and basic structure."""
    def __init__(self):
        self.skill_categories: Dict[str, List[str]] = {}
        self._load_skill_categories()
        self.section_patterns = {
            "contact": [r"contact", r"personal", r"information"],
            "summary": [r"summary", r"profile", r"objective", r"about"],
            "experience": [r"experience", r"employment", r"work", r"career", r"professional"],
            "education": [r"education", r"academic", r"qualification", r"degree"],
            "skills": [r"skills", r"technical", r"competenc", r"expertise"],
            "projects": [r"projects", r"portfolio", r"achievements"],
            "certifications": [r"certification", r"license", r"credential"],
        }

    def _load_skill_categories(self):
        try:
            from config import SKILL_CATEGORIES
            self.skill_categories = SKILL_CATEGORIES
        except Exception:
            self.skill_categories = {
                "generic": ["project management", "civil engineering", "water", "energy"],
            }

    def extract_text_from_pdf(self, file_path: str) -> str:
        try:
            doc = fitz.open(file_path)
            text = ""
            for page in doc:
                page_text = page.get_text("text")
                page_text = re.sub(r"\n\s*\n", "\n\n", page_text)
                page_text = re.sub(r"\s+", " ", page_text)
                text += page_text + "\n"
            doc.close()
            return text.strip()
        except Exception as e:
            st.error(f"PDF extraction error: {e}")
            return ""

    def extract_text_from_docx(self, file_path: str) -> str:
        try:
            d = docx.Document(file_path)
            text = ""
            for p in d.paragraphs:
                if p.text.strip():
                    text += p.text.strip() + "\n"
            for table in d.tables:
                for row in table.rows:
                    row_txt = [c.text.strip() for c in row.cells if c.text.strip()]
                    if row_txt:
                        text += " | ".join(row_txt) + "\n"
            return text.strip()
        except Exception as e:
            st.error(f"DOCX extraction error: {e}")
            return ""

    def extract_text(self, file_path: str) -> str:
        p = Path(file_path)
        ext = p.suffix.lower()
        text = ""
        if ext == ".pdf":
            text = self.extract_text_from_pdf(str(p))
        elif ext in [".docx", ".doc"]:
            text = self.extract_text_from_docx(str(p))
        elif ext == ".txt":
            try:
                with open(p, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
            except Exception as e:
                st.error(f"Text file error: {e}")
                return ""
        else:
            st.error(f"Unsupported file format: {ext}")
            return ""
        return self._clean_extracted_text(text) if text else ""

    def _clean_extracted_text(self, text: str) -> str:
        text = re.sub(r"\n\s*\n\s*\n", "\n\n", text)
        text = re.sub(r"[ \t]+", " ", text)
        return text

    def extract_contact_info(self, text: str) -> Dict[str, Any]:
        info = {"emails": [], "phones": [], "linkedin": [], "location": []}
        email_matches = re.findall(
            r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b", text, re.IGNORECASE
        )
        info["emails"] = list(set(email_matches))
        phone_matches = re.findall(
            r"(\+?[0-9]{1,3}[-.\s]?)?[0-9]{3,4}[-.\s]?[0-9]{3,4}[-.\s]?[0-9]{3,4}", text
        )
        cleaned = []
        for m in phone_matches:
            num = "".join(m)
            num_clean = re.sub(r"[^\d+]", "", num)
            if 10 <= len(num_clean) <= 15:
                cleaned.append(num_clean)
        info["phones"] = list(set(cleaned))
        linkedin_matches = re.findall(r"linkedin\.com/in/[A-Za-z0-9\-_/]+", text, re.IGNORECASE)
        info["linkedin"] = list(set(linkedin_matches))
        return info

    def extract_enhanced_skills(self, text: str) -> Dict[str, List[str]]:
        tl = text.lower()
        out: Dict[str, List[str]] = {}
        for cat, skills in self.skill_categories.items():
            found = []
            for s in skills:
                if re.search(r"\b" + re.escape(s.lower()) + r"\b", tl):
                    found.append(s)
            if found:
                out[cat] = list(sorted(set(found)))
        return out

    def extract_experience(self, text: str) -> List[Dict[str, Any]]:
        lines = [l.strip() for l in text.split("\n") if l.strip()]
        exp: List[Dict[str, Any]] = []
        for line in lines:
            m = re.match(r"(.+?)\s+at\s+(.+)", line, re.IGNORECASE)
            if m:
                pos = m.group(1).strip()
                comp = m.group(2).strip()
                exp.append({"position": pos, "company": comp, "duration": "", "description": ""})
        return exp[:8]

    def extract_education(self, text: str) -> List[Dict[str, Any]]:
        edu: List[Dict[str, Any]] = []
        matches = re.findall(
            r"(bachelor|master|phd|bsc|msc|ba|ma|mba).+?(engineering|science|arts|technology)",
            text,
            re.IGNORECASE,
        )
        for deg, field in matches:
            edu.append({"degree": deg.title(), "field": field.title(), "institution": ""})
        return edu

    def _identify_sections(self, text: str) -> Dict[str, str]:
        sections: Dict[str, str] = {}
        lines = text.split("\n")
        current = None
        buf: List[str] = []
        for line in lines:
            header = None
            for name, patterns in self.section_patterns.items():
                if any(re.search(p, line, re.IGNORECASE) for p in patterns) and len(line) < 50:
                    header = name
                    break
            if header:
                if current and buf:
                    sections[current] = "\n".join(buf)
                current, buf = header, []
            elif current:
                buf.append(line)
        if current and buf:
            sections[current] = "\n".join(buf)
        return sections

    def parse_resume(self, file_path: str, filename: str) -> Dict[str, Any]:
        try:
            raw = self.extract_text(file_path)
            if not raw.strip():
                return {"success": False, "error": "Could not extract text from resume"}
            contact = self.extract_contact_info(raw)
            skills = self.extract_enhanced_skills(raw)
            exp = self.extract_experience(raw)
            edu = self.extract_education(raw)
            sections = self._identify_sections(raw)
            word_count = len(raw.split())
            return {
                "success": True,
                "filename": filename,
                "cleaned_text": raw,
                "raw_text": raw,
                "contact_info": contact,
                "skills": skills,
                "experience": exp,
                "education": edu,
                "sections": sections,
                "statistics": {
                    "word_count": word_count,
                    "skills_found": len(skills),
                    "experience_entries": len(exp),
                    "education_entries": len(edu),
                },
                "word_count": word_count,
            }
        except Exception as e:
            st.error(f"Resume parsing error for {filename}: {e}")
            return {"success": False, "error": str(e)}

# Alias
ResumeParser = EnhancedResumeParser
